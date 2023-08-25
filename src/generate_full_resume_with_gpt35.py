import os
from docx import Document
import openai
from github import Github
import base64

# Function to read the resume from a DOCX file
def read_resume_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to read or upload a resume
def read_or_upload_resume(file_path="docs/Amitesh Tripathi Resume Aug 15.docx"):
    if os.path.exists(file_path):
        stored_resume = read_resume_from_docx(file_path)  # Use read_resume_from_docx here
        
        print("An existing resume is found. Would you like to use it? (yes/no)")
        choice = input()
        
        if choice.lower() == 'yes':
            return stored_resume
        else:
            print("Please upload your new resume (DOCX format).")
    
    # Prompt user to provide the path to the new resume
    resume_path = input("Path to your resume: ")
    
    # Read the resume from the DOCX file
    resume_data = read_resume_from_docx(resume_path)
    
    # Store it for future use
    with open(file_path, 'wb') as f:  # Change mode to 'wb'
        f.write(resume_data.encode('utf-8'))  # Encode to bytes and write
   
    return resume_data

def generate_full_resume_with_gpt35(jdtext, resumetext, api_key):
    openai.api_key = api_key
    # Generate the prompt
    prompt = f"""Based on the following job description and the existing resume data, please generate an entire resume tailored for the job. Please take inspiration from the existing resume and do not hallucinate new information. Make sure it is crafted to clear Applicant Tracking Systems (ATS) and land an interview.

    Existing Resume Data:
    {resumetext}

    Job Description:
    {jdtext}

    Please generate a full tailored resume:"""
    
    response = openai.Completion.create(
        engine="text-davinci-002",
        prompt=prompt,
        temperature=0.5,
        max_tokens=500  # You may need to adjust this based on how lengthy you expect the resume to be
    )
    
    tailored_resume = response.choices[0].text.strip()
    print(tailored_resume)
    # Ask the user for the filename
    filename = input("What filename would you like to give to the new tailored resume? ")
    
    # Ensure the filename ends with .docx
    if not filename.endswith('.docx'):
        filename += '.docx'
    
     # Define the folder and full path
    folder_path = 'data'
    full_path = os.path.join(folder_path, filename)

    # Create the new DOCX file
    doc = Document()
    doc.add_heading('Resume', 0)
    
    for line in tailored_resume.split('\n'):
        if line.strip() == '':
            continue
        elif line.strip()[-1] == ':':
            doc.add_heading(line.strip(), level=1)
        else:
            doc.add_paragraph(line.strip())
    
    # Save the new tailored resume
    doc.save(full_path)
    return full_path

'''def upload_to_github(file_path, repo_name, github_token):
    # Initialize GitHub API
    g = Github(github_token)

    # Get the repository
    repo = g.get_user().get_repo(repo_name)
    
    # Read the file and encode it into base64
    with open(file_path, "rb") as f:
        content = f.read()
    content_base64 = base64.b64encode(content)
    
    # Upload the file
    repo.create_file(
        path=file_path, 
        message="Add tailored resume", 
        content=content_base64.decode("utf-8"),
        branch="trunk"  # Replace with the branch you want to upload to
    )
    return f"https://github.com/{repo.full_name}/blob/trunk/{file_path}"
'''
def upload_to_github(file_path, repo_name, github_token):
    # Initialize GitHub API
    g = Github(github_token)

    # Get the repository
    repo = g.get_user().get_repo(repo_name)
    
    # Read the file
    with open(file_path, "rb") as f:
        content = f.read()

    # Upload the file
    repo.create_file(
        path=os.path.basename(file_path),  # Use the basename of the file to remove any directory info
        message="Add tailored resume", 
        content=content,
        branch="trunk"  # Replace with the branch you want to upload to
    )

    return f"https://github.com/{repo.full_name}/blob/trunk/{os.path.basename(file_path)}"



